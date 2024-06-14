import math
import os
import random

import docx
import convertapi
import qrcode
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, Mm
from docxtpl import DocxTemplate, InlineImage
from datetime import datetime
from dotenv import load_dotenv
from logging import basicConfig, INFO, info, error
from re import sub
from transliterate import translit
from PIL import Image, ImageDraw, ImageFont

from app import convert_date

basicConfig(
    level=INFO,
    format="%(asctime)s %(levelname)s %(message)s"
)
current_directory = os.path.dirname(os.path.abspath(__file__))
parent_directory = os.path.dirname(current_directory)

load_dotenv()

groups_dict = {
    "children": "Дети",
    "admins": "Администрация",
    "mentors": "Воспитатели",
    "teachers": "Преподаватели"
}


def insert_metrics(filepath, group_num):
    doc = DocxTemplate(f"{current_directory}/generated/{filepath}.docx")
    context = {"group_num": group_num,
               "create_time": datetime.now().strftime("%d.%m.%Y в %H:%M"),
               }
    doc.render(context)
    doc.save(f"{current_directory}/generated/{filepath}.docx")


def images_to_pdf(filepath_list: list, output_filename: str):
    images_list = []

    main_img = Image.open(filepath_list[0])
    for filepath in filepath_list[1:]:
        img = Image.open(filepath)
        img.convert('RGB')
        images_list.append(img)

    main_img.save(output_filename, save_all=True, append_images=images_list)


def convert_to_pdf(filepath: str):
    info(f"{filepath}: started converting to pdf")

    convertapi.api_secret = os.getenv('CONVERT_SECRET')

    docx_file = f"{current_directory}/generated/{filepath}.docx"

    converted = convertapi.convert(
        'pdf',
        {
            'File': docx_file
        },
        from_format='doc'
    )
    info(f"{filepath}: converted successfully!")

    converted.save_files(f"{current_directory}/generated")
    info(f"{filepath}: file save in /generated")

    if os.path.exists(docx_file):
        os.remove(docx_file)


def fill_badge(badge_type: str, name: str, caption: str):
    translitted_name = sub(r'[^a-zA-Zа-яА-ЯёЁ]', '', translit(''.join(name.split()[:2]), language_code='ru', reversed=True))  # + str(random.randint(11111, 99999))

    filename = f"badge_{badge_type}_{translitted_name}"
    info(f'Badging: generating {badge_type} for {name} ({caption})')

    initials = name.split()
    multiline = True if len(initials) == 3 else False

    fontpath = f'{current_directory}/fonts/Bebas.ttf'

    name_font = ImageFont.truetype(font=fontpath, size=85)
    caption_font = ImageFont.truetype(font=fontpath, size=60)

    badge = Image.open(f'{current_directory}/templates/badges/{badge_type}.png')
    drawer = ImageDraw.Draw(badge)

    caption_position = (badge.size[0] // 2, 530)

    drawer.text(caption_position, caption.upper(), font=caption_font, fill='black', anchor='mm', align='center')

    if multiline:
        name = f"{initials[1]} {initials[2]}\n{initials[0]}"
        drawer.multiline_text((badge.size[0] // 2, badge.size[1] // 2 + 20), name.upper(), font=name_font, fill='black', anchor='mm', align='center', spacing=15)
    else:
        drawer.text((badge.size[0] // 2, badge.size[1] // 2 + 20), name.upper(), font=name_font, fill='black', anchor='mm', align='center')

    badge.save(f"{current_directory}/generated/{filename}.png")
    # badge.show()
    info(f'Badging: saved {filename}.png')


def create_badge_sheet(final_file_name, insert_paths, spacing_mm=2, border_width_mm=0.3):
    A4_width_mm = 210
    A4_height_mm = 297

    insert_width_mm = 90
    insert_height_mm = 55

    margin_mm = 12.7

    dpi = 300
    mm_to_inches = 1 / 25.4
    A4_width_px = int(A4_width_mm * dpi * mm_to_inches)
    A4_height_px = int(A4_height_mm * dpi * mm_to_inches)
    insert_width_px = int(insert_width_mm * dpi * mm_to_inches)
    insert_height_px = int(insert_height_mm * dpi * mm_to_inches)
    margin_px = int(margin_mm * dpi * mm_to_inches)
    spacing_px = int(spacing_mm * dpi * mm_to_inches)
    border_width_px = int(border_width_mm * dpi * mm_to_inches)

    a4_image = Image.new('RGB', (A4_width_px, A4_height_px), 'white')
    draw = ImageDraw.Draw(a4_image)

    total_insert_width_px = insert_width_px + 2 * border_width_px + spacing_px
    total_insert_height_px = insert_height_px + 2 * border_width_px + spacing_px

    columns = (A4_width_px - 2 * margin_px + spacing_px) // total_insert_width_px
    rows = (A4_height_px - 2 * margin_px + spacing_px) // total_insert_height_px

    current_insert_index = 0

    for row in range(rows):
        for col in range(columns):
            if current_insert_index < len(insert_paths):
                insert_image = Image.open(f"{current_directory}/generated/{insert_paths[current_insert_index]}")
                insert_image_resized = insert_image.resize((insert_width_px, insert_height_px))

                x = margin_px + col * total_insert_width_px
                y = margin_px + row * total_insert_height_px

                draw.rectangle([x, y, x + insert_width_px + 2 * border_width_px, y + insert_height_px + 2 * border_width_px], outline="black", width=border_width_px)

                a4_image.paste(insert_image_resized, (x + border_width_px, y + border_width_px))

                current_insert_index += 1
            else:
                break

    a4_image.save(final_file_name)


def get_grouplist(group_list: list, group_num: int):
    filename = f"grouplist_{group_num}_{datetime.now().date().strftime('%d%m%Y')}"
    info(f"{filename}: creating file")

    data = []
    for child in group_list:
        child['birth'] = convert_date(str(child['birth']))
        data.append(
            [
                child['name'],
                child['birth'],
                child['comment'],
                child['parrent_name'],
                child['parrent_phone'],
            ]
        )

    doc = docx.Document(f'{current_directory}/templates/grouplist.docx')
    table = doc.tables[0]

    for i in range(len(data)):
        row = table.add_row()

        for j, value in enumerate(data[i]):
            cell = row.cells[j]
            cell_text = cell.paragraphs[0]
            run = cell_text.add_run(str(value))
            font = run.font
            font.name = "Montserrat Medium"
            if j != 2:  # не выравнивается ячейка с особенностями
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.save(f"{current_directory}/generated/{filename}.docx")
    info(f"{filename}: saved in /generated")

    insert_metrics(filename, group_num)
    info(f"{filename}: metrics inserted")

    convert_to_pdf(filename)

    return filename


def get_qr_list(group: str, group_list: list, value: str = ""):
    if value:
        extra = f"_{value}"
    else:
        extra = ""
    filename = f"qrlist_{group}{extra}_{datetime.now().date().strftime('%d%m%Y')}"
    doc = docx.Document(f'{current_directory}/templates/qr_table.docx')

    group_list.sort(key=lambda user: user['name'])
    for user in group_list:
        # print(f"creating qr {user['pass_phrase']}")
        qr_img = qrcode.make(f"https://t.me/{os.getenv('BOT_NAME')}?start={group}_{user['pass_phrase']}")
        qr_img.save(f"{current_directory}/qr/{user['pass_phrase']}.png")
        # print('OK')
    users_count = len(group_list)
    rows, cols = math.ceil(users_count / 4), 4
    qr_table = doc.add_table(rows, cols)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    index = 0
    for row in range(rows):
        for col in range(cols):
            if index + 1 <= users_count:
                cell = qr_table.cell(row, col)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                name = " ".join(group_list[index]['name'].split(" ")[:2])
                paragraph.add_run().add_text(name)
                paragraph.runs[-1].font.size = Pt(11)
                paragraph.runs[-1].font.name = "Montserrat Medium"

                image_path = f"{current_directory}/qr/{group_list[index]['pass_phrase']}.png"
                paragraph.add_run().add_picture(image_path, width=Inches(1.7))
                index += 1

    doc.save(f"{current_directory}/generated/{filename}.docx")
    if value:
        insert_metrics(filename, f"№{value}")
    else:
        insert_metrics(filename, groups_dict[group])
    convert_to_pdf(filename)

    return filename


def get_feedback(module_name: str, feedback_list: []):
    creation_date = datetime.now().date().strftime('%d_%m_%Y')
    filename = f"feedback_{module_name}_{creation_date}"
    doc = docx.Document(f'{current_directory}/templates/feedback.docx')

    title = f"Обратная связь по модулю «{module_name}» за {creation_date}"
    main_text = ""

    for fb in feedback_list:
        main_text += f"Оценка: {fb['mark']}\nКомментарий: {fb['comment']}\n\n"

    paragraph_title = doc.add_paragraph()
    run = paragraph_title.add_run(title)
    font = run.font
    font.size = Pt(15)
    font.name = "Arial"
    font.bold = True

    paragraph_body = doc.add_paragraph()
    run = paragraph_body.add_run(main_text)
    font = run.font
    font.size = Pt(13)
    font.name = "Arial"

    doc.save(f"{current_directory}/generated/{filename}.docx")
    convert_to_pdf(filename)

    return filename


def get_modules_navigation(modules_list: list, title: str):
    filename = f"navigation_{datetime.now().date().strftime('%d%m%Y')}"
    doc = docx.Document(f'{current_directory}/templates/navigation.docx')

    title_table = doc.tables[0]
    navigation_table = doc.tables[1]
    title_cell = title_table.cell(0, 1)
    paragraph_title = title_cell.paragraphs[0]
    run = paragraph_title.add_run(title)
    run.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    font = run.font
    font.name = "Montserrat SemiBold"
    font.size = Pt(20)

    for index, module in enumerate(modules_list):
        cell = navigation_table.cell(index, 0)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_text(module['name'].upper())
        paragraph.runs[-1].font.size = Pt(20)
        paragraph.runs[-1].font.name = "Montserrat Black"
        paragraph.runs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        cell = navigation_table.cell(index, 1)
        paragraph = cell.paragraphs[0]
        paragraph.add_run().add_text(module['location'].upper())
        paragraph.runs[-1].font.size = Pt(20)
        paragraph.runs[-1].font.name = "Montserrat Black"

        if index != len(modules_list) - 1:
            navigation_table.add_row()

    doc.save(f"{current_directory}/generated/{filename}.docx")
    convert_to_pdf(filename)

    return filename


def get_module_parts(children_list: list, module_info: list, teacher_info: list):
    filename = f"module_{module_info['id']}_{datetime.now().date().strftime('%d%m%Y')}"
    doc = docx.Document(f'{current_directory}/templates/module_parts.docx')

    text_title = f"{module_info['name']}" \
                 f"\n{teacher_info['name']}" \
                 f"\n{module_info['location']}"

    text_list = ""

    for index, child in enumerate(children_list):
        text_list += f"{index + 1}. {child['name']} (группа №{child['group_num']})\n"

    paragraph_title = doc.add_paragraph()
    run = paragraph_title.add_run(text_title)
    font = run.font
    font.size = Pt(15)
    font.name = "Montserrat SemiBold"
    font.bold = True

    paragraph_body = doc.add_paragraph()
    run = paragraph_body.add_run(text_list)
    font = run.font
    font.size = Pt(13)
    font.name = "Montserrat Medium"
    font.bold = False

    doc.save(f"{current_directory}/generated/{filename}.docx")
    convert_to_pdf(filename)

    return filename
